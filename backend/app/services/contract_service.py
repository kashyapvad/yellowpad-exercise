# Business logic for contract parsing and clause insertion will go here. 

import uuid
from typing import Optional
from docx import Document
import io
import re
import openai
import json


LLM_INSERTION_PLAN_PROMPT = '''
You are a contract clause placement assistant. Given a contract structure (as a list of sections), an instruction, and a clause, return a JSON object describing exactly where and how to insert the clause. Your response should be a single JSON object with these fields:
- targetType: the type of the section to target (e.g., 'heading', 'clause', 'paragraph')
- targetNumber: the number or label of the target section (e.g., '1', 'A', '11', or null if not applicable)
- targetContent: the text of the target section (if relevant)
- position: one of 'after', 'before', 'in', 'under', 'within', 'between', or 'withinContent' (for insertion within a section's content)
- childType: the type for the new clause (e.g., 'heading', 'clause', 'paragraph')
- childNumber: the number/label for the new clause (if specified or inferred)
- style: an object with any style info (e.g., bold, italics, underline, font, headingLevel)
- contentPosition: if inserting within content, specify where (e.g., 'after first sentence', 'after last paragraph', etc.)
- any other info needed for precise placement
Respond ONLY with valid JSON. Do NOT include any explanation, markdown, or text outside the JSON. Do NOT wrap your response in triple backticks or say "here is your JSON". Output ONLY the raw JSON object as shown above.

Contract structure (JSON):
"""
{contract_json}
"""
Instruction:
"""
{instruction}
"""
Clause:
"""
{clause}
"""
'''

def extract_section_number(text: str) -> Optional[str]:
    match = re.match(r"^\s*([0-9]+[A-Za-z]?|[A-Z]|[IVXLCDM]+)[.\t\s]+", text)
    return match.group(1) if match else None

def strip_section_number(text: str) -> str:
    return re.sub(r"^\s*([0-9]+[A-Za-z]?|[A-Z]|[IVXLCDM]+)[.\t\s]+", "", text).strip()

def is_all_caps(text: str) -> bool:
    letters = ''.join([c for c in text if c.isalpha()])
    return len(letters) >= 3 and letters.isupper()

def extract_contract_text(doc: Document) -> str:
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def safe_parse_llm_response(llm_content: str):
    match = re.search(r'"updated_contract"\s*:\s*"((?:[^"\\]|\\.)*)"', llm_content, re.DOTALL)
    if not match:
        raise ValueError("Could not find updated_contract in LLM response")
    value = match.group(1)
    value_fixed = re.sub(r'(?<!\\)"', r'\\"', value)
    fixed_json = '{"updated_contract": "%s"}' % value_fixed
    return json.loads(fixed_json)

def parse_docx_service(contents: bytes, instruction: Optional[str], clause: Optional[str]):
    doc = Document(io.BytesIO(contents))
    contract_text = extract_contract_text(doc)
    if instruction and clause:
        LLM_INSERT_PROMPT = '''
        You are a contract editor. Given the following contract text, and this instruction and clause, return the updated contract as a single string of the full contract text, with the clause inserted in the correct place.
        - Instruction: {instruction}
        - Clause: {clause}
        Respond ONLY with valid JSON in the following format:
        {{
          "updated_contract": "Full contract text here, with all headings, clauses, and the new clause inserted in the correct place."
        }}
        IMPORTANT:
        - Escape all double quotes in the contract text with a backslash (\").
        - Do NOT include any explanation, markdown, code block, or text outside the JSON.
        - Do NOT wrap your response in triple backticks or say "here is your JSON".
        - Output ONLY the raw JSON object as shown above.
        - If you add any extra text, unescaped quotes, or incomplete JSON, the system will break.
        - Ensure the JSON is complete and valid.
        Contract text:
        """
        {contract_text}
        """
        '''
        prompt = LLM_INSERT_PROMPT.format(
            instruction=instruction,
            clause=clause,
            contract_text=contract_text
        )
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=4096,
            response_format={"type": "json_object"},
        )
        llm_content = response.choices[0].message.content
        llm_content = llm_content.strip().replace('```json', '').replace('```', '').replace('`', '')
        result = safe_parse_llm_response(llm_content)
        updated_contract = result["updated_contract"]
        return {"updated_contract": updated_contract}
    else:
        return {"contract_text": contract_text}

def parse_contract_text_service(text: str):
    doc_stream = io.BytesIO()
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(doc_stream)
    doc_stream.seek(0)
    parsed_doc = Document(doc_stream)
    lines = [p.text.strip() for p in parsed_doc.paragraphs if p.text.strip()]
    sections = []
    for line in lines:
        type_ = 'clause'
        level = 1
        number = None
        match = re.match(r"^([0-9]+[A-Za-z]?|[A-Z])\.?\s+", line)
        if match:
            type_ = 'heading'
            number = match.group(1)
            line = re.sub(r"^([0-9]+[A-Za-z]?|[A-Z])\.?\s+", '', line).strip()
        sections.append({
            "id": str(uuid.uuid4()),
            "type": type_,
            "content": line,
            "level": level,
            "number": number,
        })
    contract = {
        "id": str(uuid.uuid4()),
        "title": "Pasted Contract",
        "sections": sections,
    }
    return contract

def parse_instruction_service(contract_json, instruction, clause):
    LLM_INSERTION_PROMPT = '''
    You are a contract clause placement assistant. Given a contract structure (as a list of sections), an instruction, and a clause, return a JSON object describing exactly where and how to insert the clause. Your response should be a single JSON object with these fields:
    - targetType: the type of the section to target (e.g., 'heading', 'clause', 'paragraph')
    - targetNumber: the number or label of the target section (e.g., '1', 'A', '11', or null if not applicable)
    - targetContent: the text of the target section (if relevant)
    - position: one of 'after', 'before', 'in', 'under', 'within', 'between', or 'withinContent' (for insertion within a section's content)
    - childType: the type for the new clause (e.g., 'heading', 'clause', 'paragraph')
    - childNumber: the number/label for the new clause (if specified or inferred)
    - style: an object with any style info (e.g., bold, italics, underline, font, headingLevel)
    - contentPosition: if inserting within content, specify where (e.g., 'after first sentence', 'after last paragraph', etc.)
    - any other info needed for precise placement
    Respond ONLY with valid JSON. Do NOT include any explanation, markdown, or text outside the JSON. Do NOT wrap your response in triple backticks or say "here is your JSON". Output ONLY the raw JSON object as shown above.

    Contract structure (JSON):
    """
    {contract_json}
    """
    Instruction:
    """
    {instruction}
    """
    Clause:
    """
    {clause}
    """
    '''
    prompt = LLM_INSERTION_PROMPT.format(
        contract_json=json.dumps(contract_json, ensure_ascii=False, indent=2),
        instruction=instruction,
        clause=clause
    )
    print("VERIFY: Sending insertion prompt to OpenAI LLM...")
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
        max_tokens=1024,
        response_format={"type": "json_object"},
    )
    llm_content = response.choices[0].message.content
    print("VERIFY: LLM raw response (insertion):", llm_content[:300])
    llm_content = llm_content.strip().replace('```json', '').replace('```', '').replace('`', '')
    try:
        result = json.loads(llm_content)
    except Exception as e:
        print("VERIFY: Error parsing LLM JSON (insertion):", e)
        raise Exception("Failed to parse LLM response as JSON (insertion).")
    return result 